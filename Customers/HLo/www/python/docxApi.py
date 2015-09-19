import traceback
import sys
import lib
import mySql
import os
from docx import Document
    
def hello(paras):
    print "Hello World"

def readFiles():
    data = []
    with open("files", "r") as fr:
        lines = fr.readlines()
        for line in lines:
            data.append(line.strip())
    return data
    
def createID(paras):
    id_list = []
    for i in range(10000):
        id_list.append(i)
    
    paras['op'] = "showItems"
    paras['internal'] = True
    data = mySql.main(paras)
    for line in data:
        num = int(line[0].split("_")[1])
        id_list.remove(num)
    
    return paras['prefix'] + "_" + str(id_list.pop())

def table_addRow(table, paras):
    row_cells = table.add_row().cells
    row_cells[0].text = paras
    
def docxInsertID(id, file):
    document = Document(file)
    table = document.tables[0]
    table_addRow(table, id)
    document.save(file)
    

def getLabel(file):
    document = Document(file)
    table = document.tables[0]
    cell = table.rows[0].cells[0]
    paragraph = cell.paragraphs[0]
    text = paragraph.text
    return text

def dbInsert(paras, id, file):
    label = getLabel(file)
    paras['op'] = "insertElement"
    paras['id'] = id
    paras['label'] = label
    paras['path'] = file
    mySql.main(paras)
    
def rename(file, id):
    dir = os.path.split(file)[0]
    file_id = dir + "\\\\" + id + ".docx"
    os.rename(file, file_id)
    return file_id
    
def insertID(paras):
    files = readFiles()
    for file in files:
        id = createID(paras)
        file = rename(file, id)
        docxInsertID(id, file)
        dbInsert(paras, id,file)
        
    
    print "status:0 , msg:Complete"

def test(paras):
    pass
    
def main():
    paras = lib.readParas()
    func = getattr(sys.modules[__name__], paras['op'])
    func(paras)
 
if __name__ == "__main__":
    main()