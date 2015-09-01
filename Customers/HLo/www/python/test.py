from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
import copy

def delete_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def insert_row(table, paras):
    row_cells = table.add_row().cells
    row_cells[0].text = str(1)
    row_cells[1].text = str(1)
    
    
document = Document('Original.docx')
allShapes = document.inline_shapes
aaa = None
for share in allShapes:
    if share.type == WD_INLINE_SHAPE.PICTURE:
        print('Shape is an embedded picture')
    elif shape.type == WD_INLINE_SHAPE.LINKED_PICTURE:
        print('Shape is a picture, but actual image file is not in this package')
    else:
        print('Shape is not a picture, got %s' % shape.type)
        '''
        inline = share._inline
        rId = inline.xpath('./a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip/@r:embed')[0]
        print rId
        image_part = document.related_parts[rId]
        image_bytes = image_part.blob
        image_stream = BytesIO(image_bytes)
        '''
        
table = document.tables[0]
cell = table.rows[0].cells[0]
paragraph = cell.paragraphs[0]
text = paragraph.text
print text

delete_row(table, table.rows[len(table.rows) - 1])
delete_row(table, table.rows[len(table.rows) - 1])

cell = table.rows[1].cells[0]
paragraph = cell.paragraphs[4]
for item in cell.paragraphs:
    print item.text

for item in table.rows:
    print item


document.save('Output_test.docx')  

document = Document()
table2 = document.add_table(rows=1, cols=2)
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'item'
hdr_cells[1].text = 'subject'
document.save('Output_test_v2.docx') 

source_document = Document('Output_test.docx')

'''

inline = shape._inline
    rId = inline.xpath('./a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip/@r:embed')[0]
    image_part = document.related_parts[rId]
    image_bytes = image_part.blob

    # write the image bytes to a file (or BytesIO stream) and feed it to document.add_picture(), maybe:
    image_stream = BytesIO(image_bytes)
    other_document.add_picture(image_stream)

    # image part also has some handy items
    image_filename = image_part.filename





docOut = Document()
docOut.add_picture(allShapes[0])
docOut.save('Output_test.docx')
'''

'''
with open('foobar.docx') as f:
    source_stream = StringIO(f.read())
document = Document(source_stream)
source_stream.close()
...
target_stream = StringIO()
document.save(target_stream)
'''